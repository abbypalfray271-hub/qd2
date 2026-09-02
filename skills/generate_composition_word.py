import os
import sys
import json
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

sys.stdout.reconfigure(encoding='utf-8')

DB_PATH = r'E:\中考库\中考语文\wiki\写作专项\中考必备七大类主题作文_结构化数据库.json'
OUTPUT_DOCX = r'E:\中考库\中考语文\outputs\中考必备七大类主题作文_精编教研版.docx'

with open(DB_PATH, 'r', encoding='utf-8') as f:
    db = json.load(f)

doc = Document()

# Page setup: Standard A4 with 2.54cm margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

# Color Palette
COLOR_NAVY = RGBColor(26, 54, 93)       # #1A365D Primary Dark Blue
COLOR_TEAL = RGBColor(44, 122, 123)     # #2C7A7B Secondary Teal
COLOR_DARK = RGBColor(45, 55, 72)       # #2D3748 Body Text
COLOR_GRAY = RGBColor(113, 128, 150)    # #718096 Subtle
COLOR_RED = RGBColor(197, 48, 48)       # #C53030 Highlights

def set_font(run, font_name="宋体", size_pt=10.5, bold=False, italic=False, color=COLOR_DARK):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color

def add_callout_box(doc, text_lines, title="【名师精析与点评】", bg_color="F7FAFC", border_color="CBD5E0"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    
    cell = table.cell(0, 0)
    # Set background shading and borders
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)
    
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="none"/>'
        f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>'
        f'  <w:bottom w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    cell._tc.get_or_add_tcPr().append(tcBorders)
    
    # Title
    p_title = cell.paragraphs[0]
    p_title.paragraph_format.space_before = Pt(4)
    p_title.paragraph_format.space_after = Pt(4)
    r_title = p_title.add_run(title)
    set_font(r_title, font_name="黑体", size_pt=10, bold=True, color=COLOR_TEAL)
    
    # Content
    for line in text_lines:
        if not line.strip():
            continue
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.2
        r = p.add_run(line.strip())
        set_font(r, font_name="楷体", size_pt=9.5, italic=False, color=COLOR_DARK)
        
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(2)
    p_spacer.paragraph_format.space_after = Pt(4)

print("Building Cover Page...")
# Cover Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(36)
p_title.paragraph_format.space_after = Pt(12)
r_main_title = p_title.add_run("中考必备七大类主题作文")
set_font(r_main_title, font_name="黑体", size_pt=26, bold=True, color=COLOR_NAVY)

# Subtitle
p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_before = Pt(0)
p_sub.paragraph_format.space_after = Pt(18)
r_sub = p_sub.add_run("—— 中考语文满分写作专项 · 核心母题与范文精编讲义")
set_font(r_sub, font_name="楷体", size_pt=14, color=COLOR_TEAL)

# Meta info
p_meta = doc.add_paragraph()
p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_meta.paragraph_format.space_before = Pt(6)
p_meta.paragraph_format.space_after = Pt(24)
r_meta = p_meta.add_run("适用学段：初中语文中考备考 · 全学科知识库工程输出物")
set_font(r_meta, font_name="宋体", size_pt=10, color=COLOR_GRAY)

# Divider Table
divider = doc.add_table(rows=1, cols=1)
divider.alignment = WD_TABLE_ALIGNMENT.CENTER
divider.columns[0].width = Inches(6.5)
d_cell = divider.cell(0, 0)
d_borders = parse_xml(
    f'<w:tcBorders {nsdecls("w")}>'
    f'  <w:bottom w:val="single" w:sz="18" w:space="0" w:color="1A365D"/>'
    f'</w:tcBorders>'
)
d_cell._tc.get_or_add_tcPr().append(d_borders)

# User Guide Summary
p_guide_head = doc.add_paragraph()
p_guide_head.paragraph_format.space_before = Pt(18)
p_guide_head.paragraph_format.space_after = Pt(6)
r_gh = p_guide_head.add_run("【中考写作通关七大母题导航】")
set_font(r_gh, font_name="黑体", size_pt=12, bold=True, color=COLOR_NAVY)

guide_table = doc.add_table(rows=8, cols=3)
guide_table.alignment = WD_TABLE_ALIGNMENT.CENTER
guide_table.columns[0].width = Inches(1.2)
guide_table.columns[1].width = Inches(1.8)
guide_table.columns[2].width = Inches(3.5)

headers = ["主题序号", "主题名称", "核心写作考向与立意锚点"]
for c_idx, h in enumerate(headers):
    cell = guide_table.cell(0, c_idx)
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E2E8F0"/>')
    cell._tc.get_or_add_tcPr().append(shd)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    set_font(r, font_name="黑体", size_pt=10, bold=True, color=COLOR_NAVY)

for r_idx, t in enumerate(db['themes']):
    row_cells = guide_table.rows[r_idx + 1].cells
    
    p0 = row_cells[0].paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(f"主题 {r_idx + 1}")
    set_font(r0, font_name="宋体", size_pt=9.5, bold=True, color=COLOR_DARK)
    
    p1 = row_cells[1].paragraphs[0]
    r1 = p1.add_run(t['theme_name'])
    set_font(r1, font_name="黑体", size_pt=9.5, color=COLOR_TEAL)
    
    p2 = row_cells[2].paragraphs[0]
    r2 = p2.add_run(t['focus'])
    set_font(r2, font_name="宋体", size_pt=9, color=COLOR_DARK)

doc.add_page_break()

# Build Each Theme
for t_idx, theme in enumerate(db['themes']):
    print(f"Adding Theme {t_idx+1}: {theme['theme_name']}...")
    
    # Theme Heading 1
    p_th = doc.add_paragraph()
    p_th.paragraph_format.space_before = Pt(12)
    p_th.paragraph_format.space_after = Pt(4)
    r_th = p_th.add_run(theme['title'])
    set_font(r_th, font_name="黑体", size_pt=18, bold=True, color=COLOR_NAVY)
    
    # Focus Badge
    p_fc = doc.add_paragraph()
    p_fc.paragraph_format.space_before = Pt(0)
    p_fc.paragraph_format.space_after = Pt(12)
    r_fc_tag = p_fc.add_run("【核心立意锚点】 ")
    set_font(r_fc_tag, font_name="黑体", size_pt=10.5, bold=True, color=COLOR_TEAL)
    r_fc_txt = p_fc.add_run(theme['focus'])
    set_font(r_fc_txt, font_name="楷体", size_pt=10.5, color=COLOR_DARK)

    # 1. 通用写作结构
    p_s1 = doc.add_paragraph()
    p_s1.paragraph_format.space_before = Pt(10)
    p_s1.paragraph_format.space_after = Pt(4)
    r_s1 = p_s1.add_run("一、通用写作结构与布局")
    set_font(r_s1, font_name="黑体", size_pt=13, bold=True, color=COLOR_NAVY)
    
    for line in theme.get('structure', []):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(line)
        set_font(r, font_name="宋体", size_pt=10.5, color=COLOR_DARK)
        
    for k, v in theme.get('special_addons', {}).items():
        p_subk = doc.add_paragraph()
        p_subk.paragraph_format.space_before = Pt(6)
        p_subk.paragraph_format.space_after = Pt(2)
        r_subk = p_subk.add_run(f"★ 特别指引：{k}")
        set_font(r_subk, font_name="黑体", size_pt=11, bold=True, color=COLOR_RED)
        for line in v:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(f"• {line}")
            set_font(r, font_name="宋体", size_pt=10, color=COLOR_DARK)

    # 2. 真题放送与审题要诀
    p_s2 = doc.add_paragraph()
    p_s2.paragraph_format.space_before = Pt(12)
    p_s2.paragraph_format.space_after = Pt(4)
    r_s2 = p_s2.add_run("二、中考真题放送与审题指导")
    set_font(r_s2, font_name="黑体", size_pt=13, bold=True, color=COLOR_NAVY)
    
    # Prompts in callout box
    prompt_lines = theme.get('exam_prompts', [])
    if prompt_lines:
        add_callout_box(doc, prompt_lines, title="【历年中考真题原题材料】", bg_color="EDF2F7", border_color="4A5568")

    p_g_head = doc.add_paragraph()
    p_g_head.paragraph_format.space_before = Pt(6)
    p_g_head.paragraph_format.space_after = Pt(2)
    r_gh = p_g_head.add_run("▶ 审题立意与避坑要点：")
    set_font(r_gh, font_name="黑体", size_pt=11, bold=True, color=COLOR_TEAL)
    for g_line in theme.get('audit_guide', []):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(f"  {g_line}")
        set_font(r, font_name="宋体", size_pt=10, color=COLOR_DARK)

    # 3. 满分标杆范文
    p_s3 = doc.add_paragraph()
    p_s3.paragraph_format.space_before = Pt(14)
    p_s3.paragraph_format.space_after = Pt(6)
    r_s3 = p_s3.add_run("三、中考满分标杆范文与逐段精析")
    set_font(r_s3, font_name="黑体", size_pt=13, bold=True, color=COLOR_NAVY)
    
    for e_idx, essay in enumerate(theme.get('exemplary_essays', [])):
        # Essay Title
        p_et = doc.add_paragraph()
        p_et.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_et.paragraph_format.space_before = Pt(10)
        p_et.paragraph_format.space_after = Pt(6)
        r_et = p_et.add_run(f"{essay['title']}")
        set_font(r_et, font_name="黑体", size_pt=12, bold=True, color=COLOR_TEAL)
        
        # Essay Paragraphs with 2-char first line indent
        for para in essay['paragraphs']:
            p_p = doc.add_paragraph()
            p_p.paragraph_format.space_before = Pt(2)
            p_p.paragraph_format.space_after = Pt(2)
            p_p.paragraph_format.line_spacing = 1.3
            p_p.paragraph_format.first_line_indent = Inches(0.28)  # ~2 chars indent
            r_p = p_p.add_run(para)
            set_font(r_p, font_name="宋体", size_pt=10.5, color=COLOR_DARK)
            
        # Analysis Box
        if essay['analysis']:
            add_callout_box(doc, essay['analysis'].split('\n'), title=f"【名师评析 · 范文{e_idx+1}亮点解析】", bg_color="F7FAFC", border_color="2C7A7B")

    # 4. 万能模板
    p_s4 = doc.add_paragraph()
    p_s4.paragraph_format.space_before = Pt(14)
    p_s4.paragraph_format.space_after = Pt(6)
    r_s4 = p_s4.add_run("四、万能模板与高分句式支架")
    set_font(r_s4, font_name="黑体", size_pt=13, bold=True, color=COLOR_NAVY)
    
    for t_line in theme.get('templates', []):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.25
        if t_line.startswith('〖') or t_line.startswith('（') or '层' in t_line[:5]:
            r = p.add_run(t_line)
            set_font(r, font_name="黑体", size_pt=10.5, bold=True, color=COLOR_TEAL)
        else:
            r = p.add_run(t_line)
            set_font(r, font_name="宋体", size_pt=10, color=COLOR_DARK)

    # 5. 名篇领航
    mp = theme.get('masterpiece', {})
    if mp and mp['title']:
        p_s5 = doc.add_paragraph()
        p_s5.paragraph_format.space_before = Pt(14)
        p_s5.paragraph_format.space_after = Pt(6)
        r_s5 = p_s5.add_run("五、名篇领航及文学鉴赏")
        set_font(r_s5, font_name="黑体", size_pt=13, bold=True, color=COLOR_NAVY)
        
        p_mpt = doc.add_paragraph()
        p_mpt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_mpt.paragraph_format.space_before = Pt(6)
        p_mpt.paragraph_format.space_after = Pt(2)
        r_mpt = p_mpt.add_run(f"《{mp['title']}》")
        set_font(r_mpt, font_name="黑体", size_pt=12, bold=True, color=COLOR_NAVY)
        
        if mp['author']:
            p_mpa = doc.add_paragraph()
            p_mpa.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_mpa.paragraph_format.space_before = Pt(0)
            p_mpa.paragraph_format.space_after = Pt(6)
            r_mpa = p_mpa.add_run(f"作者：{mp['author']}")
            set_font(r_mpa, font_name="楷体", size_pt=10, color=COLOR_GRAY)
            
        for mp_p in mp.get('content', []):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.3
            p.paragraph_format.first_line_indent = Inches(0.28)
            r = p.add_run(mp_p)
            set_font(r, font_name="宋体", size_pt=10.5, color=COLOR_DARK)
            
        if mp.get('commentary'):
            add_callout_box(doc, mp['commentary'].split('\n'), title="【名家散文鉴赏与中考借鉴指导】", bg_color="FFFDF5", border_color="D69E2E")

    # Page break between themes
    if t_idx < len(db['themes']) - 1:
        doc.add_page_break()

# Save Document
doc.save(OUTPUT_DOCX)
print(f"Word document successfully generated: {OUTPUT_DOCX}")
